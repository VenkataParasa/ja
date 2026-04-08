from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(text, level=1, color="1F4E79"):
    h = doc.add_heading(text, level=level)
    rgb = bytes.fromhex(color)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*rgb)
    return h

def para(text, bold=False, italic=False, size=10, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p

def bullet(text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p

def sub_bullet(text, size=10):
    p = doc.add_paragraph(style='List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p

def table(headers, rows, hdr_color="1F4E79", alt="EBF3FB"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], hdr_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        bg = alt if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            set_cell_bg(cells[ci], bg)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t

def divider():
    para("─" * 80, size=8, after=3)

def spacer():
    doc.add_paragraph()

# ─────────────────────────────────────────────
# PAGE SETUP
# ─────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = Inches(1)
sec.right_margin  = Inches(1)
sec.top_margin    = Inches(0.75)
sec.bottom_margin = Inches(0.75)

# ─────────────────────────────────────────────
# COVER
# ─────────────────────────────────────────────
spacer()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("JA BizTown 3.0")
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Deployment Architecture & Codebase Strategy")
r.bold = True; r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

spacer()
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}    |    Confidential & Proprietary").italic = True
spacer()

# ─────────────────────────────────────────────
# SECTION 1: EXECUTIVE SUMMARY
# ─────────────────────────────────────────────
heading("1. Executive Summary & Key Decision", level=1)
para(
    "JA BizTown 3.0 will be delivered as a single unified codebase — a Progressive Web Application (PWA) — "
    "that runs seamlessly across student iPads, Android tablets, Chromebooks, and staff laptops. There is no "
    "requirement to build separate native iOS and Android applications."
)
para(
    "However, the single application must still be distributed through the Apple App Store and Google Play Store "
    "to satisfy JA Area MDM (Mobile Device Management) requirements, in addition to being hosted on Azure for "
    "browser-based access on laptops and Chromebooks."
)
spacer()
table(
    ["Dimension", "Decision", "Q&A Evidence"],
    [
        ["Separate iOS App",       "❌ Not required",         "Q139: 'We don't expect both a mobile and a web application'"],
        ["Separate Android App",   "❌ Not required",         "Q139: 'The software should run seamlessly on one or the other'"],
        ["Single Codebase",        "✅ Mandatory",            "Q139: One unified codebase across all device types"],
        ["App Store Distribution", "✅ Required (iOS)",       "Q74: 'Areas will need to download it from the App Store'"],
        ["Play Store Distribution","✅ Required (Android)",   "Q74: 'managed through their MDM'"],
        ["Azure Web Hosting",      "✅ Required (Laptops)",   "Q114: 'Areas who continue to use their laptops'"],
        ["Mobile Phone Support",   "❌ Explicitly excluded",  "Q114: 'The simulation will NOT be run on mobile phones'"],
        ["MDM Management",         "✅ Required",             "Q74: Both App Store and MDM distribution paths needed"],
        ["Portrait + Landscape",   "✅ Required",             "Q140: UI must function in both orientations"],
    ]
)
spacer()

# ─────────────────────────────────────────────
# SECTION 2: CODEBASE STRATEGY
# ─────────────────────────────────────────────
doc.add_page_break()
heading("2. Codebase Strategy — Progressive Web Application (PWA)", level=1)
para(
    "The recommended architecture is a Progressive Web Application (PWA) built using React. A PWA is a web application "
    "that is enhanced with modern browser APIs to behave like a native app — supporting offline capability, home-screen "
    "installation, and full-screen immersive UI on tablets — while remaining a single deployable codebase."
)

heading("2.1 Why PWA Over Native App?", level=2)
table(
    ["Criterion", "PWA", "Native iOS + Android", "Why PWA Wins Here"],
    [
        ["Codebase",         "Single",          "Two separate codebases",      "Q139 explicitly requires one codebase"],
        ["Tablet Support",   "✅ Full",          "✅ Full",                     "Equivalent experience on iPad and Android"],
        ["Laptop Support",   "✅ Full (browser)","❌ Requires separate web app", "Q114 requires laptop support without extra app"],
        ["App Store Deploy", "✅ via PWA wrapper","✅ Native",                   "Capacitor/PWABuilder packages PWA for stores"],
        ["Offline Support",  "✅ Service Workers","✅ Native storage",           "Q7/Q133 require local-first offline capability"],
        ["Dev Velocity",     "Faster",          "Slower (2 teams needed)",     "One team maintains one codebase"],
        ["Cost",             "Lower",           "~40% more effort",            "Fits within $1M budget cap (Q25)"],
        ["MDM Compatible",   "✅ Yes",           "✅ Yes",                      "Both can be distributed via MDM solutions"],
        ["Updates",          "Instant (web)",   "App store approval latency",  "Faster bug fixes post-pilot per Q9"],
    ]
)
spacer()

heading("2.2 PWA Technical Capabilities Used", level=2)
para("The following PWA capabilities make the single codebase viable across all target devices:")
bullet("Service Workers — Enable local-first offline queuing of transactions (Q7/Q133). Transactions are queued in IndexedDB and synced when connectivity is restored.")
bullet("Web App Manifest — Allows the app to be installed on iPad/Android home screens in full-screen mode, behaving identically to a native app.")
bullet("IndexedDB / Cache Storage — Persistent local storage for offline transaction queues, student session state, and simulation configuration data.")
bullet("Push Notifications — Enables admin-triggered real-time simulation events (Tornado Alerts, Town Hall) via Azure Web PubSub.")
bullet("WebRTC / WebSocket — Powers the CFO real-time screen sharing feature (Q14/Q131) across 14-20 concurrent business sessions.")
bullet("Responsive Layouts — CSS Grid + Flexbox ensure the same UI renders correctly on 7-inch tablets, 10-inch iPads, and 13-inch laptop screens (Q140 — portrait and landscape).")
bullet("MediaDevices API — Used for DJ Radio microphone capture and CEO speech-to-text recording (Q17/Q20).")
spacer()

heading("2.3 Technology Stack for Single Codebase", level=2)
table(
    ["Layer", "Technology", "Purpose"],
    [
        ["UI Framework",         "React (TypeScript)",              "Component library for all screens across all device types"],
        ["Mobile Packaging",     "Capacitor.js",                    "Wraps the PWA as a native shell for App Store + Play Store submission"],
        ["State Management",     "Redux Toolkit / Zustand",         "Global state for simulation session, balances, cart state"],
        ["Offline Storage",      "IndexedDB (via Dexie.js)",        "Local transaction queue for offline-first operation (Q7/Q133)"],
        ["Real-Time",            "Azure Web PubSub SDK",            "Screen sharing, admin events, DJ broadcasting"],
        ["API Communication",    "Axios + React Query",             "Typed API calls to backend APIM with caching + retry logic"],
        ["Styling",              "CSS Modules / Styled Components",  "Responsive layout engine for tablet + laptop viewports"],
        ["Build Tool",           "Vite",                            "Fast dev builds; PWA plugin for service worker generation"],
        ["PWA Plugin",           "vite-plugin-pwa (Workbox)",       "Service worker generation for offline caching strategy"],
        ["Testing",              "Vitest + Playwright",             "Unit + E2E tests; cross-device browser automation"],
        ["CI/CD",                "Azure DevOps Pipelines",          "Q60: Vendor granted permissions to configure CI/CD pipeline"],
    ]
)
spacer()

# ─────────────────────────────────────────────
# SECTION 3: THREE DEPLOYMENT TARGETS
# ─────────────────────────────────────────────
doc.add_page_break()
heading("3. Three Deployment Targets from One Codebase", level=1)
para(
    "From the single React PWA codebase, three distinct deployment artifacts are produced through the CI/CD pipeline, "
    "each targeting a different distribution channel with no code duplication:"
)
spacer()
table(
    ["Target", "Channel", "Audience", "Format", "Hosting"],
    [
        ["Target 1", "Apple App Store",  "iPad users (student + staff)",         "IPA  (via Capacitor + Xcode)",   "Apple CDN"],
        ["Target 2", "Google Play Store","Android tablet users (student + staff)","APK/AAB (via Capacitor + Gradle)","Google CDN"],
        ["Target 3", "Azure Web (HTTPS)","Laptop + Chromebook users",            "Static Web App (PWA in browser)","Azure Static Web Apps + Front Door CDN"],
    ]
)
spacer()
para(
    "The same compiled React bundle is the input to all three targets. Capacitor.js wraps it in a thin native WebView "
    "shell for App Store/Play Store, while the Azure Static Web App hosts the same bundle directly in the browser. "
    "There is zero divergence in functionality between the three targets."
)

heading("3.1 Target 1 — Apple App Store (iOS Tablets)", level=2)
para("Relevant Q&A: Q74, Q83, Q114, Q139, Q140")
bullet("Capacitor.js wraps the React PWA in a WKWebView native iOS shell.")
bullet("The iOS project is managed in Xcode with JA-provided Apple Developer Enterprise account (Q83).")
bullet("App is submitted to the Apple App Store for distribution to JA Area MDM-managed iPads.")
bullet("Supports iPad OS (13+); optimized for both portrait and landscape orientations (Q140).")
bullet("Capacitor plugins used: Camera (QR scanning), NFC (Phase 2), MediaDevices (DJ Mic/CEO Speech), Filesystem (offline cache).")
bullet("MDM enrollment allows JA Areas to push the app silently to all managed iPads without manual install.")
bullet("TestFlight used for internal pilot distribution to 10-15 JA Areas (Q76).")
sub_bullet("Build output: signed .IPA file")
sub_bullet("Signing: JA-provided Apple Developer certificate + provisioning profile")
sub_bullet("Review time: Typically 1-3 days for App Store review; TestFlight is immediate")
spacer()

heading("3.2 Target 2 — Google Play Store (Android Tablets)", level=2)
para("Relevant Q&A: Q74, Q83, Q112, Q114, Q139")
bullet("Capacitor.js wraps the React PWA in an Android WebView native shell (Chromium-based).")
bullet("Android project built using Gradle; signed with JA-provided keystore.")
bullet("JA provides Google Play Console account per Q83.")
bullet("Distributed to JA Areas via Google Play Managed Distribution or direct APK sideloading for MDM.")
bullet("Supports Android 8.0+ (Oreo and above) covering 5-year-old tablets per Q75.")
bullet("Internal Testing Track used for pilot JA Areas before full Play Store release.")
sub_bullet("Build output: signed .AAB (Android App Bundle) — universally supported by Play Store")
sub_bullet("Direct APK also generated for MDM sideloading where Areas lack Play Store access")
sub_bullet("Review time: Typically 1-3 days for new apps; subsequent updates ~hours")
spacer()

heading("3.3 Target 3 — Azure (Web Browser for Laptops & Chromebooks)", level=2)
para("Relevant Q&A: Q114, Q139, Q140, Q57")
bullet("The same React PWA bundle is deployed to Azure Static Web Apps — no native wrapper needed.")
bullet("Students and staff on laptops or Chromebooks open the simulation URL in a modern browser (Chrome, Edge, Safari).")
bullet("Azure Front Door + CDN ensures the app loads at edge-level speed from any US facility location.")
bullet("Service workers are active in browser context — offline caching and IndexedDB work identically to the tablet apps.")
bullet("Azure Front Door WAF provides DDoS protection and enforces HTTPS across all regions.")
bullet("No separate build process — the same Vite build output deployed via Azure DevOps pipeline (Q57/Q60).")
sub_bullet("URL: Configured subdomain under JA's existing domain (e.g., biztown.ja.org)")
sub_bullet("CDN TTL: Static assets cached at edge; API calls go direct to APIM")
sub_bullet("Browser support: Chrome 90+, Edge 90+, Safari 14+, Firefox 88+")
spacer()

# ─────────────────────────────────────────────
# SECTION 4: CI/CD PIPELINE
# ─────────────────────────────────────────────
doc.add_page_break()
heading("4. CI/CD Pipeline — One Build, Three Outputs", level=1)
para(
    "Azure DevOps hosts the CI/CD pipelines. Per Q60, vendor developers are granted permissions to configure "
    "the pipeline directly. The pipeline is structured so a single git push to the appropriate branch triggers "
    "all three deployment targets automatically."
)
spacer()
table(
    ["Stage", "Action", "Output", "Trigger"],
    [
        ["1. Code Checkout",       "Pull latest from feature/main branch",                  "Source code",              "Any PR or push"],
        ["2. Install & Type Check","npm ci + TypeScript type check",                         "Type-safe build",          "Every commit"],
        ["3. Unit Tests",          "Vitest — unit + integration tests",                      "Test report + coverage",   "Every commit (gate)"],
        ["4. Build (Vite)",        "vite build — produces optimized React bundle + SW",      "dist/ folder",             "Every commit"],
        ["5. E2E Tests",           "Playwright — cross-device browser automation",           "Pass/fail + screenshots",  "PR to main (gate)"],
        ["6a. Azure Deploy",       "Azure Static Web Apps deploy action",                    "Web app live on Azure CDN","Merge to main"],
        ["6b. Android Build",      "Capacitor copy + Gradle build + sign APK/AAB",           "Signed .AAB file",         "Release tag"],
        ["6c. iOS Build",          "Capacitor copy + Xcode archive + sign IPA",              "Signed .IPA file",         "Release tag"],
        ["7a. Play Store Upload",  "Google Play Developer API — Internal Track upload",      "Available to testers",     "Release tag"],
        ["7b. App Store Upload",   "Transporter / App Store Connect API — TestFlight push",  "Available to testers",     "Release tag"],
        ["8. Smoke Tests",         "Playwright against deployed Azure environment",           "Go/no-go signal",          "Post-deploy"],
        ["9. Notify",              "MS Teams channel + JA stakeholder email summary",         "Deployment notification",  "Post-deploy"],
    ]
)
spacer()

heading("4.1 Branch Strategy", level=2)
table(
    ["Branch", "Purpose", "Deploys To", "Protected?"],
    [
        ["feature/*",    "Developer feature work",               "Local Docker only",                  "No"],
        ["develop",      "Integration testing",                  "Dev environment (Azure)",            "Yes — PR required"],
        ["staging",      "Pre-production validation",            "Staging environment (Azure)",        "Yes — QA sign-off"],
        ["main",         "Production + App Store releases",      "Prod (Azure) + App Store + Play",   "Yes — 2 approvals"],
        ["pilot/*",      "Pilot JA Area test builds",            "TestFlight + Play Internal Track",   "Yes — QA sign-off"],
    ]
)
spacer()

heading("4.2 Three Azure Environments", level=2)
para("Per the architecture, three fully isolated Azure environments are provisioned via Terraform IaC:")
table(
    ["Environment", "Purpose", "Azure Resources", "Access"],
    [
        ["Development", "Day-to-day dev testing",         "App Services, SQL Dev tier, Cosmos Dev, Redis Basic",    "Dev team only"],
        ["Staging",     "Pre-release QA + load testing",  "AKS staging cluster, SQL General Purpose, full infra",   "Dev + QA team"],
        ["Production",  "Live simulation platform",        "AKS prod cluster, SQL Business Critical, Front Door WAF","Controlled release process"],
    ]
)
spacer()

# ─────────────────────────────────────────────
# SECTION 5: MDM DISTRIBUTION
# ─────────────────────────────────────────────
doc.add_page_break()
heading("5. MDM Distribution Strategy", level=1)
para(
    "Per Q74, JA Areas will distribute the app through both the public App Store and their own MDM solution. "
    "Q112 confirms that each local JA Area owns and manages their own devices, and some Areas have MDMs while "
    "others do not — so both pathways must be supported."
)
spacer()
table(
    ["Area Type", "iOS Tablets", "Android Tablets", "Laptops / Chromebooks"],
    [
        ["Area with MDM (e.g., Jamf, Intune, VMware WS1)",
         "MDM pushes app silently from Apple VPP + App Store via JA Enterprise account",
         "Managed Google Play pushes APK silently",
         "Azure Static Web App URL pushed as bookmark or Chromebook app via MDM policy"],
        ["Area without MDM",
         "Staff installs from App Store manually using JA-shared Apple ID or direct link",
         "Staff installs from Play Store via direct link or APK sideload",
         "Staff opens browser URL; PWA install prompt appears for home screen pinning"],
    ]
)
spacer()

heading("5.1 MDM Configuration Requirements", level=2)
bullet("App Bundle ID: com.jausa.biztown (consistent across iOS + Android + Azure)")
bullet("Single Sign-On: Azure Entra B2C device session token shared via MDM-managed keychain where supported")
bullet("Kiosk / Guided Access Mode: iOS Guided Access or Android Kiosk Mode lockable by JA Area staff for simulation days")
bullet("Auto-Update Policy: MDM pushes mandatory updates before each simulation day; no manual update needed from students")
bullet("Network Policy: MDM can push Wi-Fi credentials for facility network to all enrolled devices pre-simulation")
bullet("Content Filtering: MDM restricts device to JA BizTown app only during simulation hours (optional per Area policy)")
spacer()

# ─────────────────────────────────────────────
# SECTION 6: OFFLINE ARCHITECTURE
# ─────────────────────────────────────────────
heading("6. Offline-First Architecture (Q7 / Q133)", level=1)
para(
    "Per Q7, Q15, and Q133, the platform must support 'local-first' operation — students must be able to "
    "complete key activities (purchases, deposits, payments) even during temporary connectivity disruptions. "
    "This is implemented identically across all three deployment targets (iOS/Android/Web) using browser-standard APIs "
    "available in both Capacitor-wrapped WebViews and desktop browsers."
)

heading("6.1 Offline Flow", level=2)
table(
    ["Step", "What Happens", "Technology"],
    [
        ["1. Transaction Initiated",   "Student taps 'Pay' on tablet",                                        "React UI state update"],
        ["2. Network Check",           "Service worker / network interceptor checks connectivity",              "Workbox NetworkFirst strategy"],
        ["3a. Online Path",            "API call goes to APIM → Banking microservice immediately",              "Axios + React Query"],
        ["3b. Offline Path",           "Transaction serialized and stored in local IndexedDB queue",            "Dexie.js + Service Worker"],
        ["4. Optimistic UI",           "UI shows transaction as 'Pending' — student sees success without lag",  "Redux optimistic state"],
        ["5. Connectivity Restored",   "Service worker detects network return via Background Sync API",         "Workbox BackgroundSync"],
        ["6. Queue Drain",             "Queued transactions sent to server in FIFO order with idempotency keys","Axios + UUID idempotency"],
        ["7. Conflict Resolution",     "Server-side ledger reconciliation resolves any balance discrepancies",   "Banking microservice"],
        ["8. UI Reconcile",            "Real-time WebSocket pushes final confirmed balance to all devices",      "Azure Web PubSub"],
    ]
)
spacer()

heading("6.2 Data That Must Work Offline", level=2)
table(
    ["Data Type", "Offline Available?", "Storage Location", "Sync Direction"],
    [
        ["Simulation configuration (businesses, roles, tasks)", "✅ Yes", "Service Worker Cache (pre-fetched at session start)", "Server → Device (read-only)"],
        ["Student transaction queue (POS, billing, loan payments)", "✅ Yes", "IndexedDB (Dexie.js)", "Device → Server (on reconnect)"],
        ["Student current balance (read only during offline)", "✅ Yes", "IndexedDB (last-known snapshot)", "Server → Device"],
        ["Business account balance", "✅ Yes", "IndexedDB snapshot", "Server → Device"],
        ["UI assets, CSS, JS bundles", "✅ Yes", "Service Worker Cache (Workbox)", "Server → Device (on install)"],
        ["Real-time dashboard (teacher view)", "⚠️ Degraded", "Shows last-known data; no live updates", "Resumes on reconnect"],
        ["Admin event dispatch (Tornado Alert)", "⚠️ Queued", "Queued locally; broadcast when back online", "Device → Server"],
        ["CFO screen sharing", "❌ Not available offline", "Requires active WebSocket connection", "N/A"],
    ]
)
spacer()

# ─────────────────────────────────────────────
# SECTION 7: AZURE INFRASTRUCTURE
# ─────────────────────────────────────────────
doc.add_page_break()
heading("7. Azure Infrastructure Overview", level=1)
para(
    "All backend services, APIs, and databases are hosted on Microsoft Azure. JA's Azure subscription is used, "
    "with the vendor deploying and maintaining the infrastructure via Terraform IaC (Q57)."
)
spacer()
table(
    ["Azure Service", "Role", "Tier (Prod)", "Notes"],
    [
        ["Azure Static Web Apps",       "Hosts PWA/web bundle for laptop/browser access",     "Standard",         "Global CDN; auto SSL; preview URLs per PR"],
        ["Azure Front Door (Premium)",  "Global edge entry; WAF; CDN for all static assets",  "Premium",          "WAF rules block bad actors; routes to APIM"],
        ["Azure API Management",        "BFF API gateway; rate-limiting; payload transforms",  "Developer/S1",     "Separate products per audience (tablet vs admin)"],
        ["Azure Kubernetes Service",    "Runs all backend microservices in Docker containers", "Standard D4s_v3",  "Autoscales on CPU/memory; rolling deploy"],
        ["Azure SQL Database",          "Financial ledger; student transactions; payroll",     "Business Critical","Zone-redundant; read-replica for reporting queries"],
        ["Azure Cosmos DB",             "BizTown town config; dynamic facility layouts",       "Serverless",       "Low-latency reads; no schema migrations needed"],
        ["Azure Cache for Redis",       "Idempotency keys; session caching",                  "C1 Standard",      "Prevents duplicate transaction processing"],
        ["Azure Service Bus",           "Async messaging between microservices",              "Standard",         "Transport: banking → gamification → reporting"],
        ["Azure Web PubSub",            "Real-time WebSocket for events + screen sharing",    "Standard Unit 1",  "Up to 1,000 concurrent connections"],
        ["Microsoft Entra B2C",         "Identity; session auth; JWT tokens",                 "B1 tier",          "Custom policies for student vs staff flows"],
        ["Azure Key Vault",             "Secrets; DB connection strings; certificates",        "Standard",         "Accessed only via Managed Identities"],
        ["Azure Monitor + App Insights","Distributed tracing; alerting; dashboards",          "Standard",         "Correlation IDs trace every request end-to-end"],
        ["Azure Container Registry",    "Stores Docker images for all microservices",         "Basic",            "Images tagged per git commit SHA"],
        ["Azure DevOps",                "CI/CD pipelines; source control; boards",            "Basic",            "Q60: Vendor granted pipeline configuration access"],
    ]
)
spacer()

# ─────────────────────────────────────────────
# SECTION 8: APP STORE SUBMISSION REQUIREMENTS
# ─────────────────────────────────────────────
heading("8. App Store Submission Requirements", level=1)

heading("8.1 Apple App Store (iOS / iPadOS)", level=2)
table(
    ["Requirement", "Details"],
    [
        ["Apple Developer Account",    "JA USA provides Enterprise or Standard developer account (Q83)"],
        ["Bundle ID",                   "com.jausa.biztown — consistent cross-platform"],
        ["Minimum OS",                  "iPadOS 14.0+ (covers tablets purchased within past 5 years per Q75)"],
        ["App Category",               "Education"],
        ["Age Rating",                  "4+ (no mature content; COPPA compliant — no PII collected per Q29)"],
        ["Required Capabilities",       "Camera (QR scanning), Microphone (DJ/CEO speech), NFC (Phase 2 optional)"],
        ["Privacy Manifest",            "Required for App Store compliance — declare all data collected/stored"],
        ["COPPA Compliance Statement",  "App does not collect personal data from users under 13 (Q109/Q110)"],
        ["Screenshot Requirements",     "iPad Pro 12.9\" screenshots required; iPad 11\" optional"],
        ["TestFlight Distribution",     "Used for pilot phase (10-15 JA Areas) before public App Store release"],
        ["Review Timeline",             "Initial review: 1-3 business days; updates typically within 24 hours"],
        ["Build Tool",                  "Xcode 15+ on macOS; Capacitor generates the Xcode project automatically"],
    ]
)
spacer()

heading("8.2 Google Play Store (Android Tablets)", level=2)
table(
    ["Requirement", "Details"],
    [
        ["Play Console Account",        "JA USA provides Google Play Console account (Q83)"],
        ["Application ID",              "com.jausa.biztown — matches iOS bundle ID"],
        ["Minimum SDK",                 "Android 8.0 (API 26) — covers 5-year-old Android tablets (Q75)"],
        ["Target SDK",                  "Android 14 (API 34) — required for Play Store compliance"],
        ["App Category",               "Education"],
        ["Content Rating",              "Everyone — COPPA compliant; no PII from minors"],
        ["Required Permissions",        "CAMERA (QR), RECORD_AUDIO (DJ/CEO), NFC (Phase 2)"],
        ["Data Safety Section",         "Must declare offline storage (IndexedDB equivalent), no user data transmitted off-device"],
        ["Distribution Format",         "AAB (Android App Bundle) — required by Play Store; also APK for MDM sideload"],
        ["Internal Testing Track",      "Used for pilot distribution before public release"],
        ["Managed Google Play",         "For MDM-managed Android devices without public Play Store access"],
        ["Review Timeline",             "New apps: 1-7 days; updates: hours to 1 day"],
        ["Build Tool",                  "Capacitor generates Android project; built with Gradle in CI/CD pipeline"],
    ]
)
spacer()

# ─────────────────────────────────────────────
# SECTION 9: DEVICE + BROWSER SUPPORT MATRIX
# ─────────────────────────────────────────────
doc.add_page_break()
heading("9. Device & Browser Support Matrix", level=1)
table(
    ["Device Type", "OS", "Access Method", "Priority", "Notes"],
    [
        ["iPad (9th gen+)",                "iPadOS 14+",    "App Store App (Capacitor)",      "P1 — Primary",   "Q114: Primary student device; portrait + landscape (Q140)"],
        ["Android Tablet (Samsung, Lenovo)","Android 8.0+", "Play Store App (Capacitor)",     "P1 — Primary",   "Q114: Widely used across JA Areas; MDM-managed"],
        ["Chromebook",                      "ChromeOS 80+", "Browser (Chrome) + Azure PWA",  "P1 — Required",  "Q114: Some Areas use Chromebooks; PWA installs from Chrome"],
        ["Windows Laptop",                  "Windows 10+",  "Browser (Chrome/Edge) + Azure",  "P1 — Required",  "Q114: Staff + some student Areas; Edge 90+ / Chrome 90+"],
        ["macOS Laptop",                    "macOS 11+",    "Browser (Chrome/Safari) + Azure","P2 — Supported", "Primarily staff use; Safari 14+ / Chrome 90+"],
        ["Mobile Phone (iOS/Android)",      "Any",          "Not supported",                  "❌ Excluded",    "Q114: 'The simulation will NOT be run on mobile phones'"],
    ]
)
spacer()

heading("9.1 Screen Resolution & Responsive Breakpoints", level=2)
table(
    ["Breakpoint", "Target Device", "Resolution", "Layout Behavior"],
    [
        ["Small Tablet",  "iPad Mini, 8\" Android",   "768px wide",   "Compact navigation; single-column role UI; large touch targets"],
        ["Standard Tablet","iPad 10.2\", Galaxy Tab 10\"","1024px wide","Standard navigation; 2-column layout for business dashboards"],
        ["Large Tablet",  "iPad Pro 12.9\", Tab S8+",  "1366px wide",  "Full dashboard view; multi-panel layout for teacher/admin screens"],
        ["Laptop/Desktop","Chromebook, Windows, Mac",  "1440px+ wide", "Full desktop layout; keyboard + mouse optimized; sidebar navigation"],
    ]
)
spacer()

# ─────────────────────────────────────────────
# SECTION 10: SECURITY & COMPLIANCE IN DEPLOYMENT
# ─────────────────────────────────────────────
heading("10. Security & Compliance in Deployment", level=1)
table(
    ["Requirement", "Implementation", "Q&A Reference"],
    [
        ["HTTPS Everywhere",            "Azure Front Door enforces HTTPS redirect; HSTS headers set",             "General security posture"],
        ["Zero-Trust Network",          "All microservices communicate via mTLS inside AKS; no open ports",        "Architecture principle"],
        ["No Student PII Stored",       "Only session tokens stored; no names, emails, or DOB collected",          "Q29/Q109/Q110"],
        ["COPPA Data Anonymization",    "Scheduled job anonymizes all simulation data after 30 days",              "Q110"],
        ["Audit Logging",              "All admin actions logged 30 days; user actions logged 7 days",             "Q111"],
        ["App Store Privacy Labels",    "Apple Nutrition Labels + Google Data Safety completed at submission",      "Q83/Q109"],
        ["Certificate Management",     "Azure Key Vault manages all TLS certificates with auto-renewal",           "Architecture"],
        ["Penetration Testing",         "OWASP Top 10 penetration test performed pre-launch",                      "Layer 5 QA estimate"],
        ["MDM Device Compliance",       "MDM policy enforces OS patching; old OS blocked from connecting",         "Q112"],
        ["No PhoneGap/Cordova",         "Capacitor used (modern, actively maintained; zero legacy dependency)",    "Architecture decision"],
    ]
)
spacer()

# ─────────────────────────────────────────────
# SECTION 11: SUMMARY
# ─────────────────────────────────────────────
doc.add_page_break()
heading("11. Summary & Key Principles", level=1)
para("The following principles govern all deployment and codebase decisions for JA BizTown 3.0:", bold=True)
spacer()
bullet("ONE codebase — React PWA with TypeScript. No separate native iOS or Android apps built.")
bullet("THREE deployment targets — Apple App Store, Google Play Store, Azure Static Web App — all from one build pipeline.")
bullet("Capacitor.js bridges the PWA to native App Store packaging, handling device APIs (camera, mic, NFC) without separate codebases.")
bullet("Azure DevOps CI/CD pipeline automates all three deployment targets from a single git push, with gates at unit test, E2E test, and QA sign-off stages.")
bullet("Offline-first architecture using Service Workers + IndexedDB ensures the simulation continues during temporary WiFi drops (Q7/Q133).")
bullet("MDM distribution is supported for both App Store paths and the Azure web app, accommodating Areas with and without MDM solutions (Q74/Q112).")
bullet("Mobile phone support is explicitly excluded per Q114.")
bullet("Portrait and landscape orientations are both supported per Q140.")
bullet("All security and compliance requirements (COPPA, HTTPS, WCAG, OWASP) are applied equally across all three deployment targets.")
bullet("JA provides both Apple Developer and Google Play Console accounts (Q83); the vendor manages the submission and publishing workflow.")
spacer()

divider()
para(
    f"Confidential & Proprietary — JA BizTown 3.0 Deployment Architecture Document  |  "
    f"Generated: {datetime.date.today().strftime('%B %d, %Y')}",
    size=8, italic=True
)

# ─────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────
output = "D:\\Venkata\\JA\\Deployment Needs.docx"
doc.save(output)
print(f"Saved: {output}")
