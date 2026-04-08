from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return h

def add_para(doc, text, bold=False, italic=False, size=10, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_table(doc, headers, rows, header_color="1F4E79", alt_color="EBF3FB"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_color)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg = alt_color if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            set_cell_bg(cells[c_idx], bg)
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    return table

def page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────
# PAGE SETTINGS
# ─────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(0.75)
section.bottom_margin = Inches(0.75)

# ─────────────────────────────────────────────
# COVER / TITLE
# ─────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("JA BizTown 3.0 Redevelopment")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Comprehensive Project Estimation & Analysis Report")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}    |    Rate: $85/hr (blended)    |    Budget Cap: $1,000,000").italic = True

doc.add_paragraph()

# ─────────────────────────────────────────────
# SECTION 1: EXECUTIVE SUMMARY
# ─────────────────────────────────────────────
add_heading(doc, "1. Executive Summary", level=1, color="1F4E79")

add_para(doc, (
    "This report presents the comprehensive project estimation and analysis for the redevelopment of the JA BizTown 3.0 "
    "simulation platform. The analysis is grounded in a full review of the JA BizTown RFP, all 143 vendor Question & Answer "
    "responses, and a proposed Azure-native microservices architecture designed to support 10,000–15,000 concurrent users "
    "across 60–70 simultaneous simulations nationwide."
))

add_para(doc, (
    "Two distinct estimate models are presented in this document:"
))

summary_rows = [
    ["Revised Estimate",  "$748,340", "6,310 dev hrs", "Phase 1 + Phase 2", "$850K", "Full scope within $750K budget target"],
    ["Realistic Estimate", "$995,775", "8,120 dev hrs", "Phase 1 + Phase 2 + 10% Contingency", "$1,000,000", "Most accurate full-scope with explicit risk buffer"],
]
add_table(doc, ["Estimate Model", "Total Cost", "Dev Hours", "Scope", "Budget Target", "Key Distinction"], summary_rows)
doc.add_paragraph()

add_para(doc, (
    "Both estimates are designed within JA's confirmed $1,000,000 budget cap (Q25 response) at a blended rate of $85 per hour. "
    "Each estimate includes Year 1 and Year 2 post-launch support, PM/delivery management overhead, and all explicit Q&A-driven "
    "scope requirements identified during the 143-question vendor clarification process."
))

# ─────────────────────────────────────────────
# SECTION 2: KEY Q&A FINDINGS
# ─────────────────────────────────────────────
page_break(doc)
add_heading(doc, "2. Key Q&A Findings That Drive Estimates", level=1, color="1F4E79")
add_para(doc, (
    "A thorough review of all 143 vendor Q&A responses yielded the following critical findings that directly impact "
    "scope, architecture, and effort estimates:"
))

qa_rows = [
    ["Q7, Q15, Q133", "Offline Local-First Required",    "Phase 1 Mandatory", "System must support local-first operation with sync. No full reversion to paper."],
    ["Q16",           "Business Scenarios = Grant Req.", "Phase 1 Mandatory", "Character education (ethical dilemmas, XP consequences) must be included for grant funding."],
    ["Q13, Q135",     "Canva Marketing Tool",            "Phase 1 Mandatory", "Template-based tool; step-by-step mission; image library; publish to town display."],
    ["Q14, Q131",     "CFO Screen Sharing",              "Phase 1 Mandatory", "Real-time screen mirror to 14-20 business sessions simultaneously; no audio."],
    ["Q17, Q45",      "DJ Radio Dashboard",              "Phase 1 Mandatory", "Built-in DJ UI; playlist + mic; adapts to diverse A/V systems across 98 Areas."],
    ["Q137",          "Physical Debit Card Readers",     "Phase 1 Mandatory", "3-digit account number swiped into POS. Current production model — must continue."],
    ["Q110",          "Data Anonymization (30-day)",     "Phase 1 Mandatory", "COPPA compliance: student data anonymized 30 days post-simulation."],
    ["Q68",           "Spanish Localization",            "Phase 1 Mandatory", "Vendor expected to provide all translated content."],
    ["Q77, Q78",      "Pilot On-Site Support",           "Phase 1 Required",  "Vendor must be on-site for select pilot sessions; instant availability for rest."],
    ["Q25, Q26",      "Budget Cap: $1M",                 "Hard Constraint",   "Total project cost must not exceed $1,000,000 for initial release."],
    ["Q125, Q126",    "10-15K Concurrent Users",         "Hard Constraint",   "60-70 simultaneous simulations; 50-100 TPS during peak shopping periods."],
    ["Q143",          "Full State Recovery (DR)",        "Hard Constraint",   "Simulations MUST resume from exact prior state. Restarting is not acceptable."],
    ["Q136",          "AI Features (Slogan + Speech)",   "Phase 2 Optional",  "Flexible for pilot; can be introduced post-launch if timeline is tight."],
    ["Q19, Q138",     "NFC Tap-to-Pay",                  "Phase 2 Optional",  "Nice-to-have; graceful QR code degradation required. Cross-platform iOS + Android."],
    ["Q134",          "Paper Fallback",                  "Not Required",      "Not required if offline digital mode is functional."],
    ["Q23",           "Shared Devices",                  "Removed",           "One device per student. Shared device model has been eliminated."],
    ["Q71",           "Student Accounts",                "Simplified",        "Session-based per device; no individual student accounts or PII collected."],
    ["Q57",           "Azure Deployment",                "Collaborative",     "JA provides Azure environment; vendor owns pipeline config and permissions."],
]
add_table(doc, ["Q&A Ref", "Finding", "Impact", "Details"], qa_rows)

# ─────────────────────────────────────────────
# SECTION 3: ARCHITECTURE OVERVIEW
# ─────────────────────────────────────────────
page_break(doc)
add_heading(doc, "3. Azure-Native Microservices Architecture", level=1, color="1F4E79")
add_para(doc, (
    "The proposed architecture is built on a 5-layer Azure-native microservices model. The system is designed around "
    "event-driven choreography, idempotent transaction processing, and zero-trust security principles to ensure the simulation "
    "survives massive concurrency spikes and unstable facility WiFi environments."
))

arch_rows = [
    ["Layer 1", "Edge & UI Delivery",              "Azure Front Door, CDN, Static Web Apps",                         "Student PWA, Teacher Dashboard, Admin CMS, DJ Radio, Marketing Tool"],
    ["Layer 2", "API Gateway & Security",          "Azure APIM (BFF), Entra B2C, Key Vault",                         "Auth, JWT tokens, RBAC, rate-limiting, payload transformation"],
    ["Layer 3", "Compute (Microservices)",         "AKS (Docker containers), Azure SQL, Cosmos DB, Redis Cache",      "Banking, Offline Sync, Gamification, Reporting, Voting, Scenarios"],
    ["Layer 4", "Async Event-Bus",                 "Azure Service Bus",                                               "Decouples banking from gamification; guarantees at-least-once delivery"],
    ["Layer 5", "Observability & Resilience",      "Azure Monitor, App Insights, Polly Circuit Breaker",              "Correlation ID tracing, distributed telemetry, graceful degradation"],
]
add_table(doc, ["Layer", "Name", "Azure Technologies", "Key Responsibilities"], arch_rows)
doc.add_paragraph()

add_heading(doc, "3.1 Idempotency & Transaction Safety", level=2)
add_para(doc, (
    "Because facility WiFi is unreliable, student tablets may inadvertently retry payment requests multiple times. "
    "The Transaction & Ledger Service accepts a unique Idempotency-Key (UUID) on every mutation request. On a duplicate retry, "
    "the service checks Azure Cache for Redis, recognizes the key, and returns the same successful response without "
    "double-charging the student's ledger. This pattern is enforced across all financial microservices."
))

add_heading(doc, "3.2 Offline Local-First Architecture", level=2)
add_para(doc, (
    "Per Q7/Q133, the system must support local-first operation with eventual sync. Student devices queue all "
    "transactions locally when offline. Upon reconnect, the queue syncs sequentially with a conflict reconciliation engine "
    "on the server side — ensuring every student's balance and business ledger reaches consistency without data loss. "
    "No on-premise server is required; this is a cloud-hybrid resilience pattern."
))

add_heading(doc, "3.3 Scale Requirements", level=2)
add_para(doc, (
    "Per Q125/Q126, the system must support 10,000–15,000 concurrent users across 60–70 simultaneous simulations. "
    "Peak throughput reaches 50–100 transactions per second during shopping periods. AKS autoscaling, read-replica SQL, "
    "and Azure Service Bus topic partitioning are the primary mechanisms used to achieve this scale without degradation."
))

# ─────────────────────────────────────────────
# SECTION 4: REVISED ESTIMATE ($750K)
# ─────────────────────────────────────────────
page_break(doc)
add_heading(doc, "4. Revised Estimate — $750K Budget Target", level=1, color="1F4E79")
add_para(doc, (
    "The Revised Estimate targets a total project cost of approximately $750,000 at $85/hr blended. It includes all "
    "Phase 1 (Core MVP) features and all Phase 2 (Optional) features, with the hours calibrated precisely to land within "
    "the target budget. All Q&A-mandatory items are included in Phase 1."
))

rev_summary = [
    ["Dev (Phase 1 + Phase 2)", "6,310 hrs", "$536,350"],
    ["PM / Delivery Mgmt (11%)", "694 hrs", "$58,990"],
    ["Year 1 Support (1 FTE)", "1,200 hrs", "$102,000"],
    ["Year 2 Support (0.5 FTE)", "600 hrs", "$51,000"],
    ["TOTAL", "8,804 hrs", "$748,340"],
]
add_table(doc, ["Category", "Hours", "Cost"], rev_summary)
doc.add_paragraph()

add_heading(doc, "4.1 Phase 1 — Core MVP Highlights", level=2)
p1_rows = [
    ["Foundation",         "340 hrs",   "Requirements, Architecture (10-15K scale), CI/CD, Docker Dev Setup, Data Anonymization"],
    ["UX / UI",            "1,870 hrs", "Full design system, PWA app, 6 student role UIs, CMS inheritance, Canva tool, CFO sharing, Radio, Voting"],
    ["API Gateway",         "130 hrs",  "APIM BFF, Entra B2C session auth"],
    ["Backend Compute",    "1,480 hrs", "Banking (50-100 TPS), Offline Sync, Config, Business Scenarios, Debit Card, XP, Reporting, Audit, DR"],
    ["Event Bus",            "70 hrs",  "Azure Service Bus with dead-letter handling"],
    ["Testing & Security", "1,080 hrs", "Automated tests, WCAG, Pen testing, Load testing, Cross-device, UAT"],
    ["Delivery",            "310 hrs",  "App Store, User guides, Spanish localization, Pilot support, Post-pilot sprint"],
]
add_table(doc, ["Area", "Hours", "Description"], p1_rows)
doc.add_paragraph()

add_heading(doc, "4.2 Phase 2 — Optional Features", level=2)
p2_rows = [
    ["AI Slogan Generation", "120 hrs", "Scaffolded Azure OpenAI; 3-strike guardrail; age-appropriate constrained outputs"],
    ["AI Speech-to-Text", "80 hrs", "CEO speech transcription and in-simulation display (not AI speech writing)"],
    ["NFC Tap-to-Pay", "160 hrs", "iOS + Android NFC with graceful QR fallback across device types"],
    ["Advanced Gamification", "140 hrs", "Inventory depletion; supply chain math beyond unlimited stock assumption"],
    ["4th QA Environment", "40 hrs", "Isolated regression + load test cluster separate from Dev/Staging/Prod"],
]
add_table(doc, ["Feature", "Hours", "Notes"], p2_rows)
doc.add_paragraph()

# ─────────────────────────────────────────────
# SECTION 5: REALISTIC ESTIMATE (~$996K)
# ─────────────────────────────────────────────
page_break(doc)
add_heading(doc, "5. Realistic Estimate — Full Scope with 10% Contingency", level=1, color="1F4E79")
add_para(doc, (
    "The Realistic Estimate represents the most accurate full-scope projection, incorporating all Phase 1 and Phase 2 "
    "features, with more thorough and precise hourly allocations for high-complexity items. A 10% contingency buffer is "
    "explicitly included as a named line item — not hidden — to account for scope creep, integration surprises, and unknowns "
    "that invariably emerge during a project of this complexity."
))

real_summary = [
    ["Dev (Phase 1 + Phase 2)",    "8,120 hrs", "$690,200"],
    ["Contingency Buffer (10%)",     "812 hrs",  "$69,020"],
    ["PM / Delivery Mgmt (11%)",     "983 hrs",  "$83,555"],
    ["Year 1 Support (1 FTE)",     "1,200 hrs", "$102,000"],
    ["Year 2 Support (0.5 FTE)",     "600 hrs",  "$51,000"],
    ["GRAND TOTAL",               "11,715 hrs", "$995,775"],
]
add_table(doc, ["Category", "Hours", "Cost"], real_summary)
doc.add_paragraph()

add_para(doc, (
    "The Realistic Estimate comes in at $995,775 — just under the JA-confirmed $1,000,000 budget cap. "
    "The additional hours vs the Revised Estimate reflect more thorough allocations for genuinely high-risk items "
    "including the offline sync engine, real-time screen sharing, full DR state recovery, comprehensive QA cycles, "
    "and pilot on-site support logistics."
), bold=False)

add_heading(doc, "5.1 Key Differences vs Revised Estimate", level=2)
diff_rows = [
    ["Offline Sync Engine",       "220 hrs", "260 hrs", "More complex reconciliation logic; edge case handling"],
    ["Transaction & Ledger",      "320 hrs", "360 hrs", "More thorough idempotency + audit coverage"],
    ["Performance Engineering",   "200 hrs", "220 hrs", "Additional simulation load scenarios"],
    ["Post-Pilot Sprint",         "280 hrs", "280 hrs", "Same — already conservative"],
    ["Automated Test Suite",      "280 hrs", "280 hrs", "Same"],
    ["DR & State Recovery",       "120 hrs", "130 hrs", "Full walkthrough validation included"],
    ["Voting Service (Backend)",   "70 hrs",  "80 hrs",  "Separate microservice for civic ballot processing"],
    ["Payroll Engine",             "70 hrs",  "70 hrs",  "Explicitly itemized (was bundled in banking)"],
    ["Team Building / Onboarding", "50 hrs",  "60 hrs",  "Guided first-time walkthroughs per Q42"],
    ["Post-Launch Hypercare",      "60 hrs",  "60 hrs",  "First-week instant-availability support Q78"],
    ["10% Contingency",            "N/A",     "812 hrs", "Explicit risk buffer — not in Revised"],
]
add_table(doc, ["Feature", "Revised Hrs", "Realistic Hrs", "Rationale"], diff_rows, header_color="2E75B6")
doc.add_paragraph()

# ─────────────────────────────────────────────
# SECTION 6: COMPARISON
# ─────────────────────────────────────────────
page_break(doc)
add_heading(doc, "6. Side-by-Side Comparison", level=1, color="1F4E79")

comp_rows = [
    ["Total Dev Hours",       "6,310",     "8,120"],
    ["Contingency",           "Not included", "812 hrs (10%)"],
    ["PM Overhead",           "694 hrs (11%)", "983 hrs (11%)"],
    ["Year 1 Support",        "1,200 hrs", "1,200 hrs"],
    ["Year 2 Support",        "600 hrs",   "600 hrs"],
    ["Total Hours",           "8,804 hrs", "11,715 hrs"],
    ["Total Cost (@ $85/hr)", "$748,340",  "$995,775"],
    ["Budget Headroom",       "$251,660",  "$4,225"],
    ["Phase 1 Coverage",      "Full",      "Full"],
    ["Phase 2 Coverage",      "Full",      "Full"],
    ["Contingency Buffer",    "None",      "10% explicit"],
    ["Voting Module",         "Included",  "Separately itemized"],
    ["Payroll Engine",        "Bundled",   "Separately itemized"],
    ["Hypercare (Week 1)",    "Included",  "Separately itemized"],
    ["Risk Level",            "Medium",    "Low (buffered)"],
]
add_table(doc, ["Metric", "Revised Estimate", "Realistic Estimate"], comp_rows, header_color="1F4E79")
doc.add_paragraph()

add_heading(doc, "6.1 Recommendation", level=2)
add_para(doc, (
    "For a project of this complexity — 100+ JA Area facilities, 10-15K concurrent users, real-time financial "
    "transactions, hardware integrations, and a hard January 2027 pilot deadline — the Realistic Estimate ($995,775) "
    "is the recommended baseline for client presentation. The explicit 10% contingency protects both the vendor "
    "and JA from mid-project scope discussions, and the more thorough hourly allocations reflect what a quality "
    "delivery actually requires."
), bold=False)

add_para(doc, (
    "The Revised Estimate ($748,340) remains a valuable option if JA wishes to preserve maximum budget headroom "
    "for future feature phases, hardware procurement, or Year 3 onwards support. Both models fall within the "
    "confirmed $1M budget cap."
), bold=False)

# ─────────────────────────────────────────────
# SECTION 7: TECHNOLOGY STACK
# ─────────────────────────────────────────────
page_break(doc)
add_heading(doc, "7. Technology Stack Summary", level=1, color="1F4E79")

tech_rows = [
    ["React / React Native",           "UI / Frontend",          "Cross-platform PWA for student tablets and admin SPAs on laptops"],
    ["Azure Front Door & CDN",          "Edge & Delivery",        "Global edge entry point; WAF; DDoS protection; static asset caching"],
    ["Azure Static Web Apps",           "Edge & Delivery",        "Serverless hosting for all frontends; built-in CI/CD; free SSL"],
    ["Azure API Management (APIM)",     "API Gateway",            "BFF pattern; rate-limiting; tablet vs desktop payload transforms"],
    ["Microsoft Entra B2C",             "Security/Identity",      "Session-based auth; JWT tokens; RBAC for all user types"],
    ["Azure Key Vault",                 "Security",               "All secrets managed via Managed Identities; zero-secret-in-code"],
    ["Docker",                          "Infrastructure",         "Containerizes all microservices for identical local/cloud builds"],
    ["Azure Kubernetes Service (AKS)",  "Infrastructure",         "Orchestrates and autoscales containers; zero-downtime deployments"],
    ["Azure SQL Database",              "Data Persistence",       "ACID-compliant for all financial transactions and student ledgers"],
    ["Azure Cosmos DB",                 "Data Persistence",       "Flexible JSON documents for BizTown config and facility layouts"],
    ["Azure Cache for Redis",           "Data Persistence",       "Idempotency key tracking; prevents duplicate payment processing"],
    ["Azure Web PubSub / SignalR",      "Real-Time Comms",        "WebSocket push for screen sharing, town alerts, DJ announcements"],
    ["Azure Service Bus",               "Async Messaging",        "Event-driven choreography; decouples banking from reporting/XP"],
    ["Azure Monitor & App Insights",   "Observability",           "Correlation ID tracing across all services for root-cause analysis"],
    ["Polly (Circuit Breaker)",         "Resilience",             "Retry + circuit break policies preventing cascading failures"],
    ["Microsoft Power BI",             "Analytics",               "BI layer for aggregate simulation data; export from CMS"],
    ["Azure OpenAI",                    "AI (Phase 2)",           "Managed LLM for slogan generation; speech-to-text transcription"],
]
add_table(doc, ["Technology", "Layer", "Role in JA BizTown"], tech_rows)

# ─────────────────────────────────────────────
# SECTION 8: TIMELINE & MILESTONES
# ─────────────────────────────────────────────
page_break(doc)
add_heading(doc, "8. Project Timeline & Key Milestones", level=1, color="1F4E79")

timeline_rows = [
    ["May 2026",      "Project Kickoff",               "In-person kickoff; vendor volunteers at live JA BizTown simulation (Q10/Q27)"],
    ["May–Aug 2026",  "Discovery & Architecture",       "UX research, wireframes, architecture finalization, Azure environment setup"],
    ["Jun–Oct 2026",  "Design & Core Development",      "Visual design system; all Phase 1 microservices; CMS build-out"],
    ["Sep–Nov 2026",  "Integration & Testing",          "Cross-service integration; load testing; WCAG audit; security pen testing"],
    ["Nov 2026",      "Area Training & Site Prep",       "JA Area staff training; MDM distribution; pre-pilot preparation (Q96)"],
    ["Jan–Mar 2027",  "Pilot Phase (10-15 JA Areas)",    "Full-feature pilot; vendor on-site at select Areas; instant availability for rest"],
    ["Apr 2027",      "Post-Pilot Sprint",              "Remediation of pilot feedback; defect triage; enhancement evaluation"],
    ["Jun 2027",      "Initial Release",                "Full national release; all JA Areas can opt in to new platform (Q98)"],
    ["Jun 2027+",     "Year 1 Hyper-Care",              "1 FTE support; 100% uptime M-F 8AM-7PM ET; incident escalation model"],
    ["Jun 2028+",     "Year 2 Maintenance",              "Annual upgrade cycle; infrastructure refresh; minor enhancements"],
]
add_table(doc, ["Period", "Milestone", "Description"], timeline_rows)

# ─────────────────────────────────────────────
# SECTION 9: RISK REGISTER
# ─────────────────────────────────────────────
page_break(doc)
add_heading(doc, "9. Key Risks & Mitigations", level=1, color="1F4E79")

risk_rows = [
    ["Offline Sync Conflicts",   "High",   "High",   "Mature conflict resolution strategy in architecture phase; idempotency everywhere"],
    ["WiFi Instability at Facilities", "High", "Medium", "Local-first queue; degraded-mode digital operation; no paper fallback needed"],
    ["Hardware Variability (NFC, Card Readers)", "Medium", "Low", "NFC is Phase 2; card reader SDK designed for graceful fallback to manual entry"],
    ["Scope Creep from Discovery", "High", "Low", "10% contingency explicitly budgeted; JA governance process for scope changes (Q2)"],
    ["Timeline Tightness (Jan 2027 pilot)", "High", "Medium", "Aggressive but feasible with dedicated team; post-pilot sprint buffers final release"],
    ["Legacy A/V System Diversity", "Medium", "Low", "DJ dashboard designed for wide adaptability (Bluetooth to built-in speakers)"],
    ["Device OS Fragmentation", "Medium", "Low", "PWA single codebase; cross-device testing matrix in QA budget"],
    ["Third-Party Audit Delays (WCAG)", "Low", "Low", "JA has existing auditor relationship; vendor coordinates remediation"],
    ["AI Guardrail Failures", "Low", "Low", "Phase 2; 3-strike retry limit; output filtering; no open-ended AI interaction"],
]
add_table(doc, ["Risk", "Probability", "Impact", "Mitigation"], risk_rows, header_color="C00000")
doc.add_paragraph()

# ─────────────────────────────────────────────
# FOOTER NOTE
# ─────────────────────────────────────────────
add_para(doc, "————————————————————————————————————————", size=9)
add_para(doc, (
    "This document is confidential and proprietary. All estimates are based on the JA BizTown 3.0 RFP and "
    f"Q&A responses dated March 2026. Rate: $85/hr blended. Generated: {datetime.date.today().strftime('%B %d, %Y')}."
), size=8, italic=True)

# ─────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────
output_path = "D:\\Venkata\\JA\\JA BizTown - Estimation & Analysis Report.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
