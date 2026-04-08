import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_coverage_doc():
    doc = Document()
    
    # Title
    heading = doc.add_heading('JA BizTown RFP Coverage & Traceability Report', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Intro
    doc.add_paragraph("This document provides a comprehensive traceability matrix mapping the requested requirements from the JA BizTown RFP alongside our two-phase delivery strategy.")

    # Requirements Table 1: Phases
    doc.add_heading('Phase Schedule', level=1)
    
    phases = [
        ("Phase 1", "Core MVP", "Essential features required to launch a stable, functional BizTown simulation."),
        ("Phase 2", "Nice to Have / Optional Features", "Luxury, experimental, and complex hardware integrations that enhance the experience.")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase #'
    hdr_cells[1].text = 'Phase Name'
    hdr_cells[2].text = 'Description'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                
    for phase_num, phase_name, desc in phases:
        row_cells = table.add_row().cells
        row_cells[0].text = phase_num
        row_cells[1].text = phase_name
        row_cells[2].text = desc

    doc.add_paragraph("") # Space
    
    # Traceability Matrix Section
    doc.add_heading('Traceability Matrix', level=1)
    
    # Phase 1
    doc.add_heading('Phase 1: Core Features (MVP)', level=2)
    phase1_data = [
        {
            "Requirement": "1. System Architecture & Setup",
            "Notes": "Setting up 3 environments (Dev, Staging, Prod). Pragmatic cloud architecture for stability without extreme overhead.",
            "Roles": "Solution Architect, DevOps"
        },
        {
            "Requirement": "2. UX/UI Conceptualization & Prototyping",
            "Notes": "Core user flow mapping, mood boards, and ensuring WCAG 2.2 AA accessibility standards for all base components.",
            "Roles": "Designer, FE"
        },
        {
            "Requirement": "3. Backend Core Simulation & Financial Engine",
            "Notes": "Core banking protocols, secure Entra B2C auth, unlimited inventory logic, and standard transaction processing.",
            "Roles": "Backend Developer"
        },
        {
            "Requirement": "4. CMS & Multi-Tenant Administration",
            "Notes": "React UI for staff, basic localization, and static template creation for manual paper fallback.",
            "Roles": "Full Stack (FE/BE)"
        },
        {
            "Requirement": "5. Mobile Apps & Web Dashboards",
            "Notes": "Core student onboarding flow (SPA). Payments rely seamlessly on tablet APIs (camera/QR) rather than separate external hardware.",
            "Roles": "Frontend Developer"
        },
        {
            "Requirement": "6. Testing, Delivery, & App Store",
            "Notes": "E2E testing on core flows, resolving user acceptance bugs, App Store processing, and standard security auditing.",
            "Roles": "QA, Security, PM"
        }
    ]
    
    for item in phase1_data:
        doc.add_heading(item["Requirement"], level=3)
        p = doc.add_paragraph()
        p.add_run("Coverage Notes: ").bold = True
        p.add_run(item["Notes"] + "\n")
        
        p.add_run("Relevant Roles: ").bold = True
        p.add_run(item["Roles"])

    doc.add_paragraph("") # Space

    # Phase 2
    doc.add_heading('Phase 2: Nice to Have / Optional Features', level=2)
    phase2_data = [
        {
            "Requirement": "1. Cutting-Edge AI Integrations",
            "Notes": "Integrating Azure OpenAI to auto-generate business slogans and provide speech-to-text transcription.",
            "Roles": "Backend Dev (AI Focus)"
        },
        {
            "Requirement": "2. Deep Offline Synchronization Logic",
            "Notes": "Engineering true offline data-conflict resolution across the facility for extreme stability during network outages.",
            "Roles": "Backend Dev, Solution Architect"
        },
        {
            "Requirement": "3. Automated Dynamic PDF Generation",
            "Notes": "Building custom deep-coded logic to automatically generate dynamic backup paper forms on the fly.",
            "Roles": "Full Stack (FE/BE)"
        },
        {
            "Requirement": "4. Extensive External Hardware SDK Integrations",
            "Notes": "Incorporating physical external NFC tags, external ePOS debit terminals, and A/V Radio Station hooks.",
            "Roles": "Frontend / Embedded Developer"
        },
        {
            "Requirement": "5. Intricate Gamification & Supply Line Math",
            "Notes": "Replacing unlimited stock assumptions with highly granular product depletion logistics and dynamic event generation engines.",
            "Roles": "Backend Developer"
        },
        {
            "Requirement": "6. Expanded Enterprise Environments",
            "Notes": "Adding additional completely segregated hardware environments like a dedicated standalone QA stack.",
            "Roles": "DevOps"
        }
    ]
    
    for item in phase2_data:
        doc.add_heading(item["Requirement"], level=3)
        p = doc.add_paragraph()
        p.add_run("Coverage Notes: ").bold = True
        p.add_run(item["Notes"] + "\n")
        
        p.add_run("Relevant Roles: ").bold = True
        p.add_run(item["Roles"])

    output_path = 'd:\\Venkata\\JA\\RFP Coverage Report.docx'
    doc.save(output_path)
    print(f"Report saved to {output_path}")

if __name__ == "__main__":
    create_coverage_doc()
